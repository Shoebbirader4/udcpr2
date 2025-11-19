#!/usr/bin/env python3
"""
Extract ALL rules from UDCPR and Mumbai DCPR DOCX files.
This replaces mock data with actual regulations.
"""
from pathlib import Path
from docx import Document
import json
import re
from datetime import datetime

WORK_DIR = Path("udcpr_master_data")
STAGING_DIR = WORK_DIR / "staging_rules"

def clean_text(text):
    """Clean and normalize text."""
    if not text:
        return ""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_clause_number(text):
    """Extract clause number from text."""
    # Look for patterns like: 3.1.1, 4.2, 5.3.2.1, etc.
    patterns = [
        r'\b(\d+\.\d+\.?\d*\.?\d*)\b',  # 3.1.1 or 3.1
        r'Regulation\s+No\.?\s*(\d+\.\d+\.?\d*)',  # Regulation No. 3.1
        r'Rule\s+(\d+\.\d+\.?\d*)',  # Rule 3.1
        r'Clause\s+(\d+\.\d+\.?\d*)',  # Clause 3.1
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    return None

def is_heading(para):
    """Check if paragraph is a heading."""
    style = para.style.name.lower()
    return 'heading' in style or 'title' in style

def is_list_item(para):
    """Check if paragraph is a list item."""
    style = para.style.name.lower()
    return 'list' in style

def extract_rules_from_docx(file_path, jurisdiction):
    """Extract all rules from a DOCX file."""
    print(f"\n{'='*70}")
    print(f"Extracting from: {file_path.name}")
    print(f"Jurisdiction: {jurisdiction}")
    print(f"{'='*70}\n")
    
    doc = Document(file_path)
    rules = []
    
    current_chapter = None
    current_section = None
    current_rule = None
    rule_counter = 0
    
    print(f"Total paragraphs to process: {len(doc.paragraphs)}")
    print("Processing...")
    
    for i, para in enumerate(doc.paragraphs):
        text = clean_text(para.text)
        
        if not text or len(text) < 10:
            continue
        
        # Progress indicator
        if i % 500 == 0:
            print(f"  Processed {i}/{len(doc.paragraphs)} paragraphs... ({len(rules)} rules found)")
        
        # Check if it's a heading (chapter/section)
        if is_heading(para):
            if 'chapter' in text.lower():
                current_chapter = text
            elif len(text) < 200:  # Section headings are usually short
                current_section = text
            continue
        
        # Look for clause numbers
        clause_number = extract_clause_number(text)
        
        # If we have a clause number or it's a substantial paragraph
        if clause_number or (len(text) > 50 and not text.startswith('(')):
            
            # Determine if this is a new rule or continuation
            is_new_rule = False
            
            if clause_number:
                is_new_rule = True
            elif current_rule and len(text) > 100:
                # Might be a continuation, but if it's very different, treat as new
                if not text.startswith(('provided', 'however', 'note', 'explanation')):
                    is_new_rule = True
            elif not current_rule:
                is_new_rule = True
            
            if is_new_rule:
                # Save previous rule if exists
                if current_rule:
                    rules.append(current_rule)
                
                # Create new rule
                rule_counter += 1
                
                # Generate rule ID
                if clause_number:
                    rule_id = f"{jurisdiction}_{clause_number.replace('.', '_')}"
                else:
                    rule_id = f"{jurisdiction}_rule_{rule_counter}"
                
                # Extract title (first sentence or first 100 chars)
                title = text.split('.')[0] if '.' in text else text[:100]
                title = clean_text(title)
                
                current_rule = {
                    "rule_id": rule_id,
                    "title": title,
                    "jurisdiction": jurisdiction,
                    "version": "extracted_from_docx",
                    "clause_number": clause_number or f"Section {rule_counter}",
                    "clause_text": text,
                    "chapter": current_chapter,
                    "section": current_section,
                    "parsed": {
                        "type": "rule",
                        "rule_logic": None  # Will be populated by LLM later if needed
                    },
                    "examples": [],
                    "ambiguous": False,
                    "ambiguity_reason": None,
                    "source_pdf": {
                        "filename": file_path.name,
                        "page": "extracted_from_docx",
                        "text_snippet": text[:200]
                    },
                    "created_at": datetime.now().isoformat(),
                    "extraction_method": "docx_direct"
                }
            else:
                # Append to current rule
                if current_rule:
                    current_rule["clause_text"] += " " + text
    
    # Don't forget the last rule
    if current_rule:
        rules.append(current_rule)
    
    print(f"\n✓ Extracted {len(rules)} rules from {file_path.name}")
    return rules

def extract_tables_from_docx(file_path, jurisdiction):
    """Extract rules from tables in DOCX."""
    print(f"\nExtracting tables from: {file_path.name}")
    
    doc = Document(file_path)
    table_rules = []
    
    for table_idx, table in enumerate(doc.tables):
        # Skip small tables (likely formatting)
        if len(table.rows) < 2:
            continue
        
        # Try to extract structured data from table
        headers = [clean_text(cell.text) for cell in table.rows[0].cells]
        
        # Check if this looks like a regulation table
        if any(keyword in ' '.join(headers).lower() for keyword in ['regulation', 'requirement', 'provision', 'fsi', 'setback', 'parking']):
            
            for row_idx, row in enumerate(table.rows[1:], 1):
                cells = [clean_text(cell.text) for cell in row.cells]
                
                # Skip empty rows
                if not any(cells):
                    continue
                
                # Create rule from table row
                rule_id = f"{jurisdiction}_table_{table_idx}_row_{row_idx}"
                
                # Combine cells into text
                clause_text = " | ".join([f"{headers[i]}: {cells[i]}" for i in range(min(len(headers), len(cells))) if cells[i]])
                
                if len(clause_text) > 20:  # Only if substantial content
                    table_rule = {
                        "rule_id": rule_id,
                        "title": f"Table {table_idx} - Row {row_idx}",
                        "jurisdiction": jurisdiction,
                        "version": "extracted_from_docx",
                        "clause_number": f"Table-{table_idx}.{row_idx}",
                        "clause_text": clause_text,
                        "parsed": {
                            "type": "table",
                            "table_data": {
                                "headers": headers,
                                "values": cells
                            }
                        },
                        "examples": [],
                        "ambiguous": False,
                        "ambiguity_reason": None,
                        "source_pdf": {
                            "filename": file_path.name,
                            "page": f"table_{table_idx}",
                            "text_snippet": clause_text[:200]
                        },
                        "created_at": datetime.now().isoformat(),
                        "extraction_method": "docx_table"
                    }
                    table_rules.append(table_rule)
    
    print(f"✓ Extracted {len(table_rules)} rules from tables")
    return table_rules

def main():
    """Main extraction function."""
    print("="*70)
    print("UDCPR Master - DOCX Rule Extraction")
    print("="*70)
    print("\n⚠️  This will extract ALL rules from both DOCX files")
    print("   This may take 5-10 minutes and create 500+ rules")
    print("="*70)
    
    # Create staging directory
    STAGING_DIR.mkdir(parents=True, exist_ok=True)
    
    # Find DOCX files
    root = Path(".")
    docx_files = []
    
    # Look for UDCPR file
    udcpr_files = list(root.glob("*UDCPR*.docx")) + list(root.glob("*udcpr*.docx"))
    if udcpr_files:
        docx_files.append(("maharashtra_udcpr", udcpr_files[0]))
    
    # Look for Mumbai DCPR file
    mumbai_files = list(root.glob("*MUBAI*.docx")) + list(root.glob("*mumbai*.docx")) + list(root.glob("*DCPR*.docx"))
    if mumbai_files:
        docx_files.append(("mumbai_dcpr", mumbai_files[0]))
    
    if not docx_files:
        print("\n❌ No DOCX files found!")
        print("   Please ensure DOCX files are in the project root")
        return
    
    print(f"\n✓ Found {len(docx_files)} DOCX file(s) to process\n")
    
    all_rules = []
    timestamp = int(datetime.now().timestamp())
    
    # Extract from each file
    for jurisdiction, file_path in docx_files:
        # Extract paragraph rules
        rules = extract_rules_from_docx(file_path, jurisdiction)
        all_rules.extend(rules)
        
        # Extract table rules
        table_rules = extract_tables_from_docx(file_path, jurisdiction)
        all_rules.extend(table_rules)
        
        # Save to staging
        output_file = STAGING_DIR / f"{jurisdiction}_extracted_{timestamp}.json"
        combined = rules + table_rules
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(combined, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Saved {len(combined)} rules to {output_file.name}\n")
    
    # Summary
    print("="*70)
    print("✓ EXTRACTION COMPLETE!")
    print("="*70)
    print(f"\nTotal rules extracted: {len(all_rules)}")
    
    # Count by jurisdiction
    udcpr_count = sum(1 for r in all_rules if r['jurisdiction'] == 'maharashtra_udcpr')
    mumbai_count = sum(1 for r in all_rules if r['jurisdiction'] == 'mumbai_dcpr')
    
    print(f"  - UDCPR: {udcpr_count} rules")
    print(f"  - Mumbai DCPR: {mumbai_count} rules")
    
    # Count by type
    paragraph_rules = sum(1 for r in all_rules if r['parsed']['type'] == 'rule')
    table_rules = sum(1 for r in all_rules if r['parsed']['type'] == 'table')
    
    print(f"\nBy type:")
    print(f"  - Paragraph rules: {paragraph_rules}")
    print(f"  - Table rules: {table_rules}")
    
    print(f"\nFiles saved to: {STAGING_DIR}")
    
    print("\n" + "="*70)
    print("NEXT STEPS:")
    print("="*70)
    print("\n1. Review extracted rules:")
    print("   - Check staging_rules/ directory")
    print("   - Verify rule quality")
    print()
    print("2. Auto-approve all rules:")
    print("   python scripts/auto_approve_rules.py")
    print()
    print("3. Start the application:")
    print("   cd backend && npm start")
    print("   cd frontend && npm start")
    print()
    print("4. Browse rules:")
    print("   http://localhost:3000/rules")
    print()
    print("Note: You now have REAL rules from actual UDCPR/DCPR documents!")
    print("      Mock data has been replaced with extracted regulations.")

if __name__ == "__main__":
    main()

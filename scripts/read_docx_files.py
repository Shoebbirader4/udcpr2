#!/usr/bin/env python3
"""
Read and analyze DOCX files to understand structure.
"""
from pathlib import Path
import sys

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def analyze_docx(file_path):
    """Analyze a DOCX file and print structure."""
    print(f"\n{'='*70}")
    print(f"Analyzing: {file_path.name}")
    print(f"{'='*70}\n")
    
    try:
        doc = Document(file_path)
        
        # Basic stats
        print(f"Total Paragraphs: {len(doc.paragraphs)}")
        print(f"Total Tables: {len(doc.tables)}")
        
        # Analyze styles
        styles = {}
        for para in doc.paragraphs:
            style = para.style.name
            styles[style] = styles.get(style, 0) + 1
        
        print(f"\nParagraph Styles Found:")
        for style, count in sorted(styles.items(), key=lambda x: x[1], reverse=True)[:10]:
            print(f"  {style}: {count}")
        
        # Show first 20 paragraphs
        print(f"\nFirst 20 Paragraphs:")
        print("-" * 70)
        for i, para in enumerate(doc.paragraphs[:20], 1):
            text = para.text.strip()
            if text:
                style = para.style.name
                print(f"{i}. [{style}] {text[:100]}{'...' if len(text) > 100 else ''}")
        
        # Analyze tables
        if doc.tables:
            print(f"\nFirst Table Structure:")
            print("-" * 70)
            table = doc.tables[0]
            print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            print("\nFirst 3 rows:")
            for i, row in enumerate(table.rows[:3], 1):
                cells = [cell.text.strip()[:30] for cell in row.cells]
                print(f"Row {i}: {' | '.join(cells)}")
        
        # Look for clause patterns
        print(f"\nSearching for Clause Patterns:")
        print("-" * 70)
        import re
        clause_pattern = re.compile(r'\b\d+\.\d+\.?\d*\.?\d*\b')
        
        clauses_found = []
        for para in doc.paragraphs[:100]:  # Check first 100 paragraphs
            text = para.text.strip()
            if clause_pattern.search(text):
                clauses_found.append(text[:150])
        
        print(f"Found {len(clauses_found)} paragraphs with clause numbers in first 100 paragraphs")
        if clauses_found:
            print("\nSample clauses:")
            for clause in clauses_found[:5]:
                print(f"  • {clause}...")
        
        return True
        
    except Exception as e:
        print(f"Error reading {file_path.name}: {e}")
        return False

def main():
    """Main function to read DOCX files."""
    print("="*70)
    print("DOCX File Reader and Analyzer")
    print("="*70)
    
    # Look for DOCX files in project root
    root = Path(".")
    docx_files = list(root.glob("*.docx")) + list(root.glob("*.DOCX"))
    
    if not docx_files:
        print("\n❌ No DOCX files found in project root")
        print("\nPlease ensure you have:")
        print("  - UDCPR.docx (or similar)")
        print("  - Mumbai-DCPR.docx (or similar)")
        print("\nPlace them in the project root directory.")
        return
    
    print(f"\n✓ Found {len(docx_files)} DOCX file(s):")
    for f in docx_files:
        print(f"  • {f.name}")
    
    # Analyze each file
    for docx_file in docx_files:
        analyze_docx(docx_file)
    
    print("\n" + "="*70)
    print("Analysis Complete!")
    print("="*70)
    print("\nNext steps:")
    print("  1. Review the structure above")
    print("  2. Run: python scripts/extract_from_docx.py")
    print("     (This will extract all rules)")

if __name__ == "__main__":
    main()
